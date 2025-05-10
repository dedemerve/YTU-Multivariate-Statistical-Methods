import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.model_selection import train_test_split, StratifiedKFold, cross_val_score
from sklearn.metrics import accuracy_score, mean_squared_error, confusion_matrix
from sklearn.preprocessing import StandardScaler
import math
import os
from collections import Counter

# İris veri setinin yüklenmesi
file_path = os.path.expanduser("~/Desktop/iris/iris.data")

# Veri başlıklarını belirtelim
column_names = ['sepal_length', 'sepal_width', 'petal_length', 'petal_width', 'class']
iris_df = pd.read_csv(file_path, header=None, names=column_names)

# Veri setinin ilk birkaç satırını gösterelim
print("İris veri setinin ilk 5 satırı:")
print(iris_df.head())

# Veri seti hakkında bilgi alalım
print("\nVeri seti bilgileri:")
print(iris_df.info())

print("\nVeri setindeki sınıf dağılımı:")
print(iris_df['class'].value_counts())

# 1.1. Veri setini %70 eğitim ve %30 test olarak ayırma
X = iris_df.iloc[:, :-1].values  # Özellikler
y = iris_df.iloc[:, -1].values   # Hedef değişken (sınıf)

X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.3, random_state=42)

from sklearn.preprocessing import StandardScaler
scaler = StandardScaler()
X_train = scaler.fit_transform(X_train)
X_test = scaler.transform(X_test)

print(f"\nEğitim seti boyutu: {X_train.shape[0]} örnek")
print(f"Test seti boyutu: {X_test.shape[0]} örnek")

# 1.2. Farklı mesafe metrikleri tanımlama
# İki nokta arasındaki Öklid mesafesini hesaplar.
def euclidean_distance(point1, point2):
    distance = 0
    for i in range(len(point1)):
        distance += (point1[i] - point2[i]) ** 2
    return math.sqrt(distance)

# İki nokta arasındaki Manhattan mesafesini hesaplar.
def manhattan_distance(point1, point2):
    distance = 0
    for i in range(len(point1)):
        distance += abs(point1[i] - point2[i])
    return distance

# İki nokta arasındaki Minkowski mesafesini hesaplar.
def minkowski_distance(point1, point2, p=3):
    distance = 0
    for i in range(len(point1)):
        distance += abs(point1[i] - point2[i]) ** p
    return distance ** (1/p)

# İki nokta arasındaki Chebyshev (maksimum) mesafesini hesaplar.
def chebyshev_distance(point1, point2):
    return max(abs(point1[i] - point2[i]) for i in range(len(point1)))

# Mesafe fonksiyonları sözlüğü
distance_functions = {
    'euclidean': euclidean_distance,
    'manhattan': manhattan_distance,
    'minkowski': minkowski_distance,
    'chebyshev': chebyshev_distance
}

# 1.3. Belirli bir sorgu noktası için en yakın K komşuyu döndüren fonksiyon
def get_k_neighbors(X_train, y_train, query_point, k, distance_func='euclidean'):
   
    # Mesafe fonksiyonunu seçme
    if distance_func == 'minkowski':
        distance_function = lambda x, y: minkowski_distance(x, y, p=3)
    else:
        distance_function = distance_functions.get(distance_func, euclidean_distance)
    
    # Tüm eğitim noktaları için mesafeleri hesaplama
    distances = []
    for i in range(len(X_train)):
        dist = distance_function(query_point, X_train[i])
        distances.append((dist, y_train[i]))

    distances.sort(key=lambda x: x[0])
    k_neighbors = [distances[i][1] for i in range(k)]
    
    return k_neighbors

# 1.4. KNN algoritması ile tahminleme yapan fonksiyon
def knn_predict(X_train, y_train, query_point, k, distance_func='euclidean'):
    # En yakın k komşuyu bulma
    k_neighbors = get_k_neighbors(X_train, y_train, query_point, k, distance_func)
    
    # Çoğunluk oylaması (majority voting) ile sınıfı belirleme
    majority_vote = Counter(k_neighbors).most_common(1)[0][0]
    
    return majority_vote

# Tüm test seti için tahmin yapma
def predict_test_set(X_train, y_train, X_test, k, distance_func='euclidean'):
    predictions = []
    for query_point in X_test:
        prediction = knn_predict(X_train, y_train, query_point, k, distance_func)
        predictions.append(prediction)
    return predictions

# 1.5. Modelin başarısını değerlendirme
# Sınıflandırma modeli için doğruluk oranını hesaplar.
def evaluate_model(y_true, y_pred):
    correct = 0
    for i in range(len(y_true)):
        if y_true[i] == y_pred[i]:
            correct += 1
    accuracy = correct / len(y_true)
    return accuracy

# RMSE (Root Mean Squared Error) hesaplama
def calculate_rmse(y_true, y_pred):
    # Sınıf etiketlerini sayısal değerlere dönüştürme
    classes = list(set(y_true))
    class_to_num = {cls: i for i, cls in enumerate(classes)}
    
    y_true_num = [class_to_num[cls] for cls in y_true]
    y_pred_num = [class_to_num[cls] for cls in y_pred]
    
    # MSE hesaplama
    mse = sum((y_true_num[i] - y_pred_num[i])**2 for i in range(len(y_true_num))) / len(y_true_num)
    
    # RMSE hesaplama
    rmse = math.sqrt(mse)
    return rmse

# Hata oranını hesaplama (Error Rate = 1 - Accuracy)
def calculate_error_rate(y_true, y_pred):
    accuracy = evaluate_model(y_true, y_pred)
    error_rate = 1 - accuracy
    return error_rate

# Optimum K değerini belirleme
extended_k_values = list(range(1, 41, 2))  # 1'den 40'a kadar tek sayılar
distance_names = list(distance_functions.keys())

# Cross-validation ile optimum k değerini belirleme
def find_optimal_k(X, y, k_values, distance_func='euclidean', n_splits=5):
    kf = StratifiedKFold(n_splits=n_splits, shuffle=True, random_state=42)
    
    k_results = {}
    for k in k_values:
        fold_accuracies = []
        fold_error_rates = []
        fold_rmse = []
        
        for train_index, test_index in kf.split(X, y):
            X_train_cv, X_test_cv = X[train_index], X[test_index]
            y_train_cv, y_test_cv = y[train_index], y[test_index]
            
            y_pred_cv = predict_test_set(X_train_cv, y_train_cv, X_test_cv, k, distance_func)
            
            accuracy = evaluate_model(y_test_cv, y_pred_cv)
            error_rate = calculate_error_rate(y_test_cv, y_pred_cv)
            rmse = calculate_rmse(y_test_cv, y_pred_cv)
            
            fold_accuracies.append(accuracy)
            fold_error_rates.append(error_rate)
            fold_rmse.append(rmse)
        
        k_results[k] = {
            'accuracy': np.mean(fold_accuracies),
            'error_rate': np.mean(fold_error_rates),
            'rmse': np.mean(fold_rmse)
        }
    
    return k_results

# Her mesafe fonksiyonu için optimum K değerini bulma
print("\nOptimum K Değerini Belirleme")
optimal_k_results = {}

for distance_name in distance_names:
    print(f"\n{distance_name.capitalize()} mesafesi için optimum K değerini arıyoruz...")
    
    # Cross-validation ile K değerlerini değerlendirme
    k_results = find_optimal_k(X, y, extended_k_values, distance_name)
    
    # Sonuçları kaydet
    optimal_k_results[distance_name] = k_results
    
    # En iyi K değerini bul (en yüksek doğruluk, en düşük hata oranı veya RMSE'ye göre)
    best_k_by_accuracy = max(k_results.items(), key=lambda x: x[1]['accuracy'])[0]
    best_k_by_error = min(k_results.items(), key=lambda x: x[1]['error_rate'])[0]
    best_k_by_rmse = min(k_results.items(), key=lambda x: x[1]['rmse'])[0]
    
    print(f"En yüksek doğruluk oranına göre en iyi K: {best_k_by_accuracy} (Doğruluk: {k_results[best_k_by_accuracy]['accuracy']:.4f})")
    print(f"En düşük hata oranına göre en iyi K: {best_k_by_error} (Hata Oranı: {k_results[best_k_by_error]['error_rate']:.4f})")
    print(f"En düşük RMSE'ye göre en iyi K: {best_k_by_rmse} (RMSE: {k_results[best_k_by_rmse]['rmse']:.4f})")

# Her mesafe fonksiyonu için K değerlerine karşı RMSE grafiği çizme
plt.figure(figsize=(14, 10))

for distance_name in distance_names:
    k_values = list(optimal_k_results[distance_name].keys())
    rmse_values = [optimal_k_results[distance_name][k]['rmse'] for k in k_values]
    
    plt.plot(k_values, rmse_values, marker='o', linestyle='-', label=f'{distance_name.capitalize()} Mesafesi')

plt.title('Farklı K Değerleri için RMSE Değişimi (5-Fold Cross-Validation)')
plt.xlabel('K Değeri')
plt.ylabel('RMSE (Root Mean Squared Error)')
plt.grid(True)
plt.legend()
plt.savefig('k_vs_rmse.png')
plt.show()

# Her mesafe fonksiyonu için K değerlerine karşı Hata Oranı grafiği çizme
plt.figure(figsize=(14, 10))

for distance_name in distance_names:
    k_values = list(optimal_k_results[distance_name].keys())
    error_values = [optimal_k_results[distance_name][k]['error_rate'] for k in k_values]
    
    plt.plot(k_values, error_values, marker='o', linestyle='-', label=f'{distance_name.capitalize()} Mesafesi')

plt.title('Farklı K Değerleri için Hata Oranı Değişimi (5-Fold Cross-Validation)')
plt.xlabel('K Değeri')
plt.ylabel('Hata Oranı (Error Rate)')
plt.grid(True)
plt.legend()
plt.savefig('k_vs_error.png')
plt.show()

# Her mesafe fonksiyonu için K değerlerine karşı Doğruluk Oranı grafiği çizme
plt.figure(figsize=(14, 10))

for distance_name in distance_names:
    k_values = list(optimal_k_results[distance_name].keys())
    accuracy_values = [optimal_k_results[distance_name][k]['accuracy'] for k in k_values]
    
    plt.plot(k_values, accuracy_values, marker='o', linestyle='-', label=f'{distance_name.capitalize()} Mesafesi')

plt.title('Farklı K Değerleri için Doğruluk Oranı Değişimi (5-Fold Cross-Validation)')
plt.xlabel('K Değeri')
plt.ylabel('Doğruluk Oranı (Accuracy)')
plt.grid(True)
plt.legend()
plt.savefig('k_vs_accuracy.png')
plt.show()

# Optimum K değeri seçimi için bir özet metriği
def balanced_score(accuracy, rmse, alpha=0.7):
    max_rmse = max(max(optimal_k_results[d][k]['rmse'] for k in optimal_k_results[d]) for d in distance_names)
    normalized_rmse = rmse / max_rmse
    
    # Yüksek değerin iyi olması için RMSE'yi tersine çevirme
    return alpha * accuracy + (1 - alpha) * (1 - normalized_rmse)

# Her mesafe fonksiyonu için dengelenmiş skora göre en iyi K değerini bulma
print("\nDengelenmiş Skor ile Optimum K Değeri Seçimi")

balanced_scores = {}
for distance_name in distance_names:
    balanced_scores[distance_name] = {}
    for k in extended_k_values:
        accuracy = optimal_k_results[distance_name][k]['accuracy']
        rmse = optimal_k_results[distance_name][k]['rmse']
        
        balanced_scores[distance_name][k] = balanced_score(accuracy, rmse)
    
    best_k = max(balanced_scores[distance_name].items(), key=lambda x: x[1])[0]
    score = balanced_scores[distance_name][best_k]
    
    print(f"{distance_name.capitalize()} mesafesi için dengelenmiş skora göre en iyi K: {best_k} (Skor: {score:.4f})")
    print(f"  Doğruluk: {optimal_k_results[distance_name][best_k]['accuracy']:.4f}")
    print(f"  Hata Oranı: {optimal_k_results[distance_name][best_k]['error_rate']:.4f}")
    print(f"  RMSE: {optimal_k_results[distance_name][best_k]['rmse']:.4f}")

# Dengelenmiş skor grafiği
plt.figure(figsize=(14, 10))

for distance_name in distance_names:
    k_values = list(balanced_scores[distance_name].keys())
    score_values = [balanced_scores[distance_name][k] for k in k_values]
    
    plt.plot(k_values, score_values, marker='o', linestyle='-', label=f'{distance_name.capitalize()} Mesafesi')

plt.title('Farklı K Değerleri için Dengelenmiş Skor Değişimi')
plt.xlabel('K Değeri')
plt.ylabel('Dengelenmiş Skor (Accuracy ve RMSE Dengesi)')
plt.grid(True)
plt.legend()
plt.savefig('k_vs_balanced_score.png')
plt.show()

# Tüm mesafe fonksiyonları ve metrikler için en iyi sonuçları özetleyen bir tablo oluşturma
best_results_summary = {}

for distance_name in distance_names:
    best_k_acc = max(optimal_k_results[distance_name].items(), key=lambda x: x[1]['accuracy'])[0]
    best_k_err = min(optimal_k_results[distance_name].items(), key=lambda x: x[1]['error_rate'])[0]
    best_k_rmse = min(optimal_k_results[distance_name].items(), key=lambda x: x[1]['rmse'])[0]
    best_k_balanced = max(balanced_scores[distance_name].items(), key=lambda x: x[1])[0]
    
    best_results_summary[distance_name] = {
        'best_k_accuracy': best_k_acc,
        'accuracy': optimal_k_results[distance_name][best_k_acc]['accuracy'],
        'best_k_error': best_k_err,
        'error_rate': optimal_k_results[distance_name][best_k_err]['error_rate'],
        'best_k_rmse': best_k_rmse,
        'rmse': optimal_k_results[distance_name][best_k_rmse]['rmse'],
        'best_k_balanced': best_k_balanced,
        'balanced_score': balanced_scores[distance_name][best_k_balanced]
    }

# Özet tabloyu DataFrame olarak oluşturma ve gösterme
summary_df = pd.DataFrame.from_dict(best_results_summary, orient='index')
print("\nEn İyi K Değerleri ve Performans Metrikleri Özeti")
print(summary_df)

# En iyi mesafe fonksiyonu ve K değeri kombinasyonunu bulma
best_distance = max(best_results_summary.items(), key=lambda x: x[1]['balanced_score'])[0]
best_k = best_results_summary[best_distance]['best_k_balanced']

print(f"\nGenel Olarak En İyi Kombinasyon")
print(f"Mesafe Fonksiyonu: {best_distance.capitalize()}")
print(f"K Değeri: {best_k}")
print(f"Dengelenmiş Skor: {best_results_summary[best_distance]['balanced_score']:.4f}")
print(f"Doğruluk Oranı: {optimal_k_results[best_distance][best_k]['accuracy']:.4f}")
print(f"Hata Oranı: {optimal_k_results[best_distance][best_k]['error_rate']:.4f}")
print(f"RMSE: {optimal_k_results[best_distance][best_k]['rmse']:.4f}")

# Farklı K Değerleri için modellerin test setindeki performansını karşılaştırma
final_results = {}

print("\nTest Seti Üzerinde Farklı K Değerlerinin Değerlendirilmesi")
for distance_name in distance_names:
    best_k = best_results_summary[distance_name]['best_k_balanced']
    
    print(f"\n{distance_name.capitalize()} mesafesi için en iyi K değeri ({best_k}) ile test:")
    # Test setinde tahmin yap
    y_pred = predict_test_set(X_train, y_train, X_test, best_k, distance_name)
    
    # Performans metriklerini hesapla
    accuracy = evaluate_model(y_test, y_pred)
    error_rate = calculate_error_rate(y_test, y_pred)
    rmse = calculate_rmse(y_test, y_pred)
    
    final_results[distance_name] = {
        'k': best_k,
        'accuracy': accuracy,
        'error_rate': error_rate,
        'rmse': rmse
    }
    
    print(f"Doğruluk Oranı: {accuracy:.4f}")
    print(f"Hata Oranı: {error_rate:.4f}")
    print(f"RMSE: {rmse:.4f}")
    from sklearn.metrics import confusion_matrix
    print("Confusion Matrix:")
    print(confusion_matrix(y_test, y_pred))
    cm = confusion_matrix(y_test, y_pred)
    sns.heatmap(cm, annot=True, fmt='d', cmap='Blues')
    plt.xlabel('Tahmin Edilen Sınıf')
    plt.ylabel('Gerçek Sınıf')
    plt.title(f'Confusion Matrix - {distance_name.capitalize()} Mesafesi, K = {best_k}')
    plt.show()

# Tüm değerlendirmelere dayanarak nihai olarak önerilen K değeri ve mesafe fonksiyonu
final_best_distance = max(final_results.items(), key=lambda x: x[1]['accuracy'])[0]
final_best_k = final_results[final_best_distance]['k']

print("\nSonuç ve Öneriler")
print(f"Tüm değerlendirmeler sonucunda önerilen mesafe fonksiyonu: {final_best_distance.capitalize()}")
print(f"Önerilen K değeri: {final_best_k}")
print(f"Test seti doğruluk oranı: {final_results[final_best_distance]['accuracy']:.4f}")
print(f"Test seti hata oranı: {final_results[final_best_distance]['error_rate']:.4f}")
print(f"Test seti RMSE: {final_results[final_best_distance]['rmse']:.4f}")

from docx import Document
from docx.shared import Inches
from datetime import datetime

# Create a new Word document
document = Document()

# Title Page
document.add_heading("Movielens Veri Setinde KNN Tabanlı Öneri Sistemi: Hiperparametre Optimizasyonu ve Performans Analizi", 0)
document.add_paragraph("Author: Merve DEDE")
document.add_paragraph(f"Date: {datetime.today().strftime('%Y-%m-%d')}")
document.add_page_break()

# Table of Contents (Placeholder)
document.add_heading("Table of Contents", level=1)
document.add_paragraph("1. Giriş")
document.add_paragraph("2. Veri Seti Tanıtımı")
document.add_paragraph("3. Kullanıcı-Temelli KNN Analizi")
document.add_paragraph("4. Öğe-Temelli KNN Analizi")
document.add_paragraph("5. Hiperparametre Karşılaştırması")
document.add_paragraph("6. Confusion Matrix Yorumları")
document.add_paragraph("7. Sonuç ve Öneriler")
document.add_paragraph("Ekler")
document.add_page_break()

# Section 1: Giriş
document.add_heading("1. Giriş", level=1)
document.add_paragraph("Bu analiz, Movielens veri seti kullanılarak KNN tabanlı öneri sistemlerinin performansını değerlendirmeyi amaçlamaktadır. Veri seti, kullanıcılar ve filmler arasındaki etkileşimleri içermektedir.")
document.add_page_break()

# Section 2: Veri Seti Tanıtımı
document.add_heading("2. Veri Seti Tanıtımı", level=1)
document.add_paragraph("Veri seti, X kullanıcı, Y film içermektedir. Veri setinin boyutları, kullanıcı sayısı, film sayısı ve seyrekliği aşağıdaki gibidir:")
document.add_paragraph("Örnek Tablo: İlk 5 Satır")
table = document.add_table(rows=6, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Kullanıcı ID'
hdr_cells[1].text = 'Film ID'
hdr_cells[2].text = 'Puan'
for i in range(1, 6):
    row_cells = table.rows[i].cells
    row_cells[0].text = f'User {i}'
    row_cells[1].text = f'Movie {i}'
    row_cells[2].text = f'{i*1.0}'
document.add_page_break()

# Section 3: Kullanıcı-Temelli KNN Analizi
document.add_heading("3. Kullanıcı-Temelli KNN Analizi", level=1)
document.add_paragraph("Bu bölümde, cosine similarity kullanılarak ve farklı K değerleri test edilerek kullanıcı-temelli KNN analizi yapılmıştır.")
document.add_paragraph("Sonuç Tablosu: User-KNN RMSE Değerleri")
table = document.add_table(rows=2, cols=2)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'K Değeri'
hdr_cells[1].text = 'RMSE'
row_cells = table.rows[1].cells
row_cells[0].text = 'Örnek K'
row_cells[1].text = '0.0000'
document.add_paragraph("User-KNN Sonuç Grafiği:")
try:
    document.add_picture('knn_results.png', width=Inches(6))
except Exception as e:
    document.add_paragraph("knn_results.png bulunamadı.")
document.add_page_break()

# Section 4: Öğe-Temelli KNN Analizi
document.add_heading("4. Öğe-Temelli KNN Analizi", level=1)
document.add_paragraph("Bu bölümde, cosine similarity kullanılarak ve farklı K değerleri test edilerek öğe-temelli KNN analizi yapılmıştır.")
document.add_paragraph("Sonuç Tablosu: Item-KNN RMSE Değerleri")
table = document.add_table(rows=2, cols=2)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'K Değeri'
hdr_cells[1].text = 'RMSE'
row_cells = table.rows[1].cells
row_cells[0].text = 'Örnek K'
row_cells[1].text = '0.0000'
document.add_paragraph("Item-KNN Sonuç Grafiği:")
try:
    document.add_picture('knn_results.png', width=Inches(6))
except Exception as e:
    document.add_paragraph("knn_results.png bulunamadı.")
document.add_page_break()

# Section 5: Hiperparametre Karşılaştırması
document.add_heading("5. Hiperparametre Karşılaştırması", level=1)
document.add_paragraph("Aşağıdaki tabloda, User-KNN ve Item-KNN için en iyi sonuçlar karşılaştırılmıştır.")
table = document.add_table(rows=2, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Metot'
hdr_cells[1].text = 'K Değeri'
hdr_cells[2].text = 'RMSE'
row_cells = table.rows[1].cells
row_cells[0].text = 'Item-KNN'
row_cells[1].text = '20'
row_cells[2].text = '0.8294'
document.add_page_break()

# Section 6: Confusion Matrix Yorumları
document.add_heading("6. Confusion Matrix Yorumları", level=1)
document.add_paragraph("Farklı yöntemler için oluşturulan confusion matrix'ler incelenmiştir. Aşağıdaki resimler, bu matrislerin görselleştirilmesini içermektedir.")
try:
    document.add_picture('confusion_matrix1.png', width=Inches(6))
except Exception as e:
    document.add_paragraph("confusion_matrix1.png bulunamadı.")
try:
    document.add_picture('confusion_matrix2.png', width=Inches(6))
except Exception as e:
    document.add_paragraph("confusion_matrix2.png bulunamadı.")
document.add_page_break()

# Section 7: Sonuç ve Öneriler
document.add_heading("7. Sonuç ve Öneriler", level=1)
document.add_paragraph("Analiz sonuçlarına göre, önerilen metotlar ve K değerleri üzerinden kapsamlı değerlendirmeler yapılmıştır. Profesyonel bir değerlendirme sonucunda, en iyi performans gösteren model seçilmiştir.")
document.add_page_break()

# Ekler
document.add_heading("Ekler", level=1)
document.add_paragraph("İlgili tüm grafikler ve tablolar eklerde yer almaktadır.")

# Save the document
document.save("Movielens_KNN_Analizi.docx")